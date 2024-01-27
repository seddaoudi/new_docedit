import os
import uuid

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from docxtpl import DocxTemplate
from docx import Document
from doc2docx import convert
import docx2txt
from fastapi.responses import FileResponse


app = FastAPI(docs_url="/api/doc", openapi_url="/api/openapi.json", redoc_url="/api/redoc")


def create_unique_folder():
    # Generate a unique folder name using uuid
    folder_name = str(uuid.uuid4())
    folder_path = os.path.join("uploaded_files", folder_name)

    # Create the unique folder
    os.makedirs(folder_path)

    return folder_path


def delete_files_except_final(folder_path):
    # Delete files in the specified folder except those starting with "final_"
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if not file_name.startswith("final_"):
            os.remove(file_path)


@app.post("/api/word_submit",
          operation_id="replaceData",
          description="this endpoint is for replacing data in a doc file or docx",
          summary="this endpoint is for replacing data in a doc file or docx",
          response_class=FileResponse,
          response_description="the response is a file"
          )
async def edit(footer_text: str, main_file: UploadFile = File(...), logo_file: UploadFile = File(...)):
    try:
        # Create a unique folder for this request
        unique_folder_path = create_unique_folder()

        # Save the uploaded main file to the unique folder
        main_file_path = os.path.join(unique_folder_path, main_file.filename)
        with open(main_file_path, "wb") as f:
            f.write(main_file.file.read())

        # Check if the uploaded file has a .doc extension and convert to .docx if needed
        if main_file_path.lower().endswith(".doc"):
            convert(main_file_path)
            main_file_path = main_file_path.replace(".doc", ".docx")

        # Modify the document using the provided footer_text
        document = Document(main_file_path)
        for section in document.sections:
            footer = section.footer
            for run in footer.paragraphs[0].runs:
                print(f"=>{run.text.encode('ascii', 'ignore').decode('ascii')}")
            footer.paragraphs[0].runs[3].text = footer_text
            footer.paragraphs[0].runs[5].text = "000000000"
            footer.paragraphs[0].runs[8].text = "QQQQQQQQ"

        # Save the modified document
        modified_file_path = os.path.join(unique_folder_path, f"modified_{main_file.filename}")
        document.save(modified_file_path)

        # Save the uploaded logo file to the unique folder
        logo_file_path = os.path.join(unique_folder_path, logo_file.filename)
        with open(logo_file_path, "wb") as f:
            f.write(logo_file.file.read())

        # Use the modified document as a template and replace media
        doc_template = DocxTemplate(modified_file_path)
        context = {}

        logos_folder = "logo"
        docx2txt.process(modified_file_path, unique_folder_path)
        doc_template.replace_media(os.path.join(unique_folder_path, "image2.png"), logo_file_path)

        doc_template.render(context)

        # Save the final document
        final_file_path = os.path.join(unique_folder_path, f"final_{main_file.filename}")
        doc_template.save(final_file_path)

        # Delete files in the unique folder except those starting with "final_"
        delete_files_except_final(unique_folder_path)

        # Return the final document as a downloadable file
        return FileResponse(final_file_path, filename=f"final_{main_file.filename}")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
